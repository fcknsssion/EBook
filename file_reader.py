import os
import re
try:
    import docx
except ImportError:
    docx = None
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
try:
    from pdf2image import convert_from_path
except ImportError:
    convert_from_path = None
try:
    import tabula
except ImportError:
    tabula = None
from utils import resource_path

file_cache = {}
image_cache = {}


def try_read_file(filepath):
    cache_key = os.path.abspath(filepath)
    if cache_key in file_cache:
        return file_cache[cache_key]["content"], file_cache[cache_key]["images"], None

    text_extensions = ('.txt', '.md', '.log')
    pdf_extensions = ('.pdf',)
    docx_extensions = ('.docx',)
    doc_extensions = ('.doc',)
    file_ext = os.path.splitext(filepath)[1].lower()
    content_with_styles = []
    images = []

    if file_ext in text_extensions:
        encodings = ['utf-8', 'windows-1251', 'cp866', 'latin1']
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as file:
                    lines = file.readlines()
                    file_dir = os.path.dirname(filepath)
                    for i, line in enumerate(lines):
                        line = line.strip()
                        if not line:
                            content_with_styles.append(("", False, False, False, False, False, "left"))
                            continue
                        is_bold = bool(re.search(r"^(Тема|Цель|Задание|Задания)\s*:?", line, re.IGNORECASE))
                        is_image = os.path.splitext(line)[1].lower() in ('.png', '.jpg', '.jpeg', '.gif')
                        if is_image:
                            image_path = os.path.join(file_dir, line)
                            if os.path.exists(image_path):
                                images.append(image_path)
                                content_with_styles.append(
                                    (f"[Изображение: {line}]", False, False, False, False, False, "left"))
                            else:
                                content_with_styles.append(
                                    (f"[Отсутствует изображение: {line}]", False, False, False, False, False, "left"))
                        else:
                            is_list = line.startswith("- ") or re.match(r"^\d+\.\s", line)
                            content_with_styles.append((line, is_list, i == 0, False, False, is_bold, "left"))
                    file_cache[cache_key] = {"content": content_with_styles, "images": images}
                    return content_with_styles, images, None
            except UnicodeDecodeError:
                continue
            except Exception as e:
                return None, [], f"Ошибка при чтении текстового файла: {e}"
        return None, [], "Не удалось определить кодировку текстового файла."

    elif file_ext in pdf_extensions:
        if PyPDF2 is None and convert_from_path is None:
            return None, [], "Библиотеки PyPDF2 и pdf2image не установлены."
        try:
            if tabula:
                tables = tabula.read_pdf(filepath, pages='all', multiple_tables=True)
                for table in tables:
                    table_data = table.values.tolist()
                    content_with_styles.append((table_data, False, False, False, 11, False, "left"))
            images = convert_pdf_to_images(filepath)
            if PyPDF2:
                with open(filepath, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(reader.pages[:5]):
                        text = page.extract_text() or ""
                        lines = text.split('\n')
                        for i, line in enumerate(lines):
                            line = line.strip()
                            if not line:
                                content_with_styles.append(("", False, False, False, False, False, "left"))
                                continue
                            is_bold = bool(re.search(r"^(Тема|Цель|Задание|Задания)\s*:?", line, re.IGNORECASE))
                            content_with_styles.append(
                                (line, False, page_num == 0 and i == 0, False, False, is_bold, "left"))
            file_cache[cache_key] = {"content": content_with_styles, "images": images}
            return content_with_styles, images, None
        except Exception as e:
            return None, [], f"Ошибка при чтении PDF файла: {e}"

    elif file_ext in docx_extensions:
        if docx is None:
            return None, [], "Библиотека python-docx не установлена."
        try:
            doc = docx.Document(filepath)
            images = extract_docx_images(filepath)
            for i, element in enumerate(doc.element.body):
                if isinstance(element, docx.oxml.table.CT_Tbl):
                    table = docx.table.Table(element, doc)
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            cell_text = ' '.join(para.text.strip() for para in cell.paragraphs if para.text.strip())
                            row_data.append(cell_text)
                        table_data.append(row_data)
                    content_with_styles.append((table_data, False, False, False, 11, False, "left"))
                elif isinstance(element, docx.oxml.text.paragraph.CT_P):
                    para = docx.text.paragraph.Paragraph(element, doc)
                    text = para.text.strip()
                    if not text:
                        content_with_styles.append(("", False, False, False, False, False, "left"))
                        continue
                    is_bold = False
                    is_italic = False
                    font_size = 11
                    alignment = "left"
                    if para.runs:
                        for run in para.runs:
                            if run.bold:
                                is_bold = True
                            if run.italic:
                                is_italic = True
                            if run.font.size:
                                font_size = run.font.size.pt
                    if para.alignment:
                        alignment_map = {0: "left", 1: "center", 2: "right", 3: "left"}
                        alignment_value = para.alignment
                        alignment = alignment_map.get(alignment_value, "left")
                    content_with_styles.append((text, is_bold, i == 0, is_italic, font_size, False, alignment))
            file_cache[cache_key] = {"content": content_with_styles, "images": images}
            return content_with_styles, images, None
        except Exception as e:
            return None, [], f"Ошибка при чтении .docx файла: {e}"

    elif file_ext in doc_extensions:
        return None, [], "Формат .doc не поддерживается. Пожалуйста, конвертируйте файл в .docx с помощью Microsoft Word или LibreOffice."

    else:
        return None, [], f"Формат {file_ext} не поддерживается."


def extract_docx_images(docx_path):
    try:
        doc = docx.Document(docx_path)
        images = []
        temp_dir = resource_path("temp_images")
        os.makedirs(temp_dir, exist_ok=True)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref.lower():
                image_ext = rel.target_ref.split('.')[-1].lower()
                if image_ext not in ('png', 'jpg', 'jpeg', 'gif'):
                    continue
                image_data = rel.target_part.blob
                image_name = f"image_{len(images)}.{image_ext}"
                image_path = os.path.join(temp_dir, image_name)
                with open(image_path, 'wb') as f:
                    f.write(image_data)
                images.append(image_path)
        return images
    except Exception:
        return []


def convert_pdf_to_images(pdf_path):
    if convert_from_path is None:
        return []
    try:
        images = convert_from_path(pdf_path, dpi=600, fmt='png')
        temp_dir = resource_path("temp_images")
        os.makedirs(temp_dir, exist_ok=True)
        image_paths = []
        for i, img in enumerate(images[:5]):
            image_path = os.path.join(temp_dir, f"pdf_page_{i}.png")
            img.save(image_path, 'PNG', quality=95)
            image_paths.append(image_path)
        return image_paths
    except Exception:
        return []